"""Genera la plantilla DOCX rellenable que el cliente descarga y completa."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "/home/claude/scorm_builder_proyecto/plantilla/Plantilla_Curso_SCORM.docx"


def add_styled_paragraph(doc, text, style=None, bold=False, italic=False,
                          color=None, size=None, alignment=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if style:
        p.style = doc.styles[style]
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_callout_paragraph(doc, prefix, text, bg_color="DBEAFE"):
    """Inserta un párrafo de tipo callout con fondo coloreado."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_color)
    pPr.append(shd)
    run = p.add_run(f"[{prefix}] ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0A2540")
    p.add_run(text)
    return p


# ========================= CREAR DOCUMENTO =========================
doc = Document()

# Configurar fuente y tamaño base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ========================= PORTADA / METADATOS =========================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PLANTILLA DE CURSO SCORM')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string('0A2540')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Rellena este documento siguiendo las instrucciones y conviértelo en un SCORM con tu marca')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string('475569')

doc.add_paragraph()

# ========================= INSTRUCCIONES =========================
add_styled_paragraph(doc, '📌 Antes de empezar', style='Heading 2')

doc.add_paragraph(
    'Esta plantilla está pensada para que tú escribas tu contenido sin preocuparte '
    'por la parte técnica. Solo tienes que respetar tres reglas:'
)
doc.add_paragraph(
    '1. Usa los estilos de Word (Título 1, Título 2, Título 3) para los encabezados, '
    'no formatees el texto a mano cambiando tamaños.',
    style='List Number'
)
doc.add_paragraph(
    '2. Mantén las palabras clave entre corchetes que ves en los ejemplos: '
    '[CLAVE], [ALERTA], [EXITO], [EJEMPLO], etc.',
    style='List Number'
)
doc.add_paragraph(
    '3. Sustituye el texto de ejemplo por tu contenido real, pero conserva la estructura.',
    style='List Number'
)

doc.add_paragraph()

# ========================= METADATOS =========================
add_styled_paragraph(doc, '⚙️ Metadatos del curso', style='Heading 2')

doc.add_paragraph(
    'Rellena estos campos. Pueden modificarse luego en la app si lo prefieres.'
)

meta_table = doc.add_table(rows=6, cols=2)
meta_table.style = 'Light Grid'
meta_table.autofit = True

meta_data = [
    ('Título del curso', 'Escribe aquí el título completo del curso'),
    ('Subtítulo / colectivo destinatario', 'Por ejemplo: Formación obligatoria · personal docente'),
    ('Autor o entidad', 'Tu nombre, asociación o consultora'),
    ('Sector', 'educación / sanitario / deportivo / corporativo / asociativo'),
    ('Paleta de color preferida', 'azul / crimson / teal / verde / morado / personalizada'),
    ('Mastery score (% mínimo aprobado)', '70'),
]
for i, (k, v) in enumerate(meta_data):
    meta_table.cell(i, 0).text = k
    meta_table.cell(i, 1).text = v
    for cell in meta_table.row_cells(i):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if cell == meta_table.cell(i, 0):
                    run.bold = True
                run.font.size = Pt(10)

doc.add_paragraph()

# ========================= EJEMPLO TEMA 1 =========================
add_styled_paragraph(doc, 'Tema 1. Sustituye este título por el de tu primer tema', style='Heading 1')

# Subapartado 1.1
add_styled_paragraph(doc, '1.1. Primer subapartado del tema', style='Heading 2')

doc.add_paragraph(
    'Aquí escribes el primer párrafo de contenido. Puedes redactar tantos párrafos '
    'como necesites. Cada párrafo separado por un salto de línea se trata como un '
    'párrafo independiente en el SCORM resultante.'
)

doc.add_paragraph(
    'Las palabras importantes puedes ponerlas en negrita normalmente '
    'usando Ctrl+B o el botón de negrita de Word. También puedes usar cursiva.'
)

# Callout clave
add_callout_paragraph(
    doc,
    'CLAVE',
    'Aquí va una idea clave que el alumno debe retener. Usa este bloque para '
    'destacar conceptos centrales del tema.',
    bg_color="DBEAFE"
)

doc.add_paragraph()

# Subapartado 1.2
add_styled_paragraph(doc, '1.2. Segundo subapartado', style='Heading 2')

doc.add_paragraph(
    'Otro párrafo de contenido. Las listas con viñetas se mantienen tal cual:'
)
doc.add_paragraph('Primer ítem de la lista', style='List Bullet')
doc.add_paragraph('Segundo ítem de la lista', style='List Bullet')
doc.add_paragraph('Tercer ítem de la lista', style='List Bullet')

doc.add_paragraph(
    'Las listas numeradas también:'
)
doc.add_paragraph('Primera acción', style='List Number')
doc.add_paragraph('Segunda acción', style='List Number')
doc.add_paragraph('Tercera acción', style='List Number')

# Diferentes tipos de callout
add_callout_paragraph(
    doc, 'ALERTA',
    'Bloque rojo para advertencias o conductas que hay que evitar.',
    bg_color="FEE2E2"
)

add_callout_paragraph(
    doc, 'EXITO',
    'Bloque verde para buenas prácticas o confirmaciones.',
    bg_color="D1FAE5"
)

add_callout_paragraph(
    doc, 'CUIDADO',
    'Bloque amarillo para precauciones o matices importantes.',
    bg_color="FEF3C7"
)

doc.add_paragraph()

# Subapartado 1.3 con tabla
add_styled_paragraph(doc, '1.3. Tablas', style='Heading 2')

doc.add_paragraph(
    'Las tablas se preservan automáticamente. Aquí un ejemplo:'
)

ej_table = doc.add_table(rows=4, cols=3)
ej_table.style = 'Light Grid Accent 1'
hdr = ej_table.rows[0].cells
hdr[0].text = 'Concepto'
hdr[1].text = 'Definición'
hdr[2].text = 'Ejemplo'
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

filas_ej = [
    ('Primer concepto', 'Su definición clara', 'Un ejemplo concreto'),
    ('Segundo concepto', 'Otra definición', 'Otro ejemplo'),
    ('Tercer concepto', 'Tercera definición', 'Tercer ejemplo'),
]
for i, fila in enumerate(filas_ej, start=1):
    for j, txt in enumerate(fila):
        ej_table.cell(i, j).text = txt

doc.add_paragraph()

# Subapartado 1.4 con ejemplo
add_styled_paragraph(doc, '1.4. Casos prácticos y ejemplos', style='Heading 2')

doc.add_paragraph(
    'Para incluir un caso práctico destacado, usa el bloque [EJEMPLO]:'
)

add_callout_paragraph(
    doc, 'EJEMPLO',
    'Título del caso práctico',
    bg_color="EFF6FF"
)
doc.add_paragraph(
    'Aquí va la descripción del caso. Puede ocupar varios párrafos. '
    'Describe la situación, los personajes implicados, el contexto.'
)

add_callout_paragraph(
    doc, 'REFLEXION',
    '¿Qué pregunta debe hacerse el alumno antes de seguir leyendo?',
    bg_color="FEF3C7"
)

add_callout_paragraph(
    doc, 'ANALISIS',
    'Análisis del caso. Esto aparecerá en un bloque desplegable que el alumno '
    'abre cuando quiera ver la respuesta. Aquí puedes ser todo lo extenso '
    'que necesites.',
    bg_color="D1FAE5"
)

doc.add_paragraph()

# Subapartado 1.5 con cita
add_styled_paragraph(doc, '1.5. Citas y referencias legales', style='Heading 2')

doc.add_paragraph(
    'Cuando necesites citar una norma, ley o autoridad, usa el bloque [CITA]:'
)

add_callout_paragraph(
    doc, 'CITA',
    'FUENTE: Artículo 16 de la Ley Orgánica 8/2021 (LOPIVI)\n'
    '"Toda persona que en el marco de su actividad profesional tenga conocimiento '
    'de una situación de violencia ejercida sobre una persona menor de edad..."',
    bg_color="F1F5F9"
)

# ========================= QUIZ =========================
doc.add_paragraph()
add_styled_paragraph(doc, '1.6. Quiz del tema', style='Heading 2')

doc.add_paragraph(
    'Escribe aquí las preguntas del test, siguiendo este formato exacto. '
    'Mínimo 4 preguntas, recomendado 8-15. Si prefieres que la app las genere '
    'automáticamente con IA, deja esta sección vacía.'
)

quiz_text = """1. ¿Cuál es la afirmación correcta sobre el concepto principal del tema?
A. Una opción incorrecta
B. La opción correcta
C. Otra opción incorrecta
D. Cuarta opción incorrecta
Correcta: B
Explicación: La opción B es correcta porque... (esta línea es opcional)

2. ¿Qué se debe hacer ante la situación X?
A. Acción inadecuada
B. Acción adecuada
C. Acción inadecuada distinta
D. Acción inadecuada cuarta
Correcta: B

3. Selecciona la respuesta verdadera:
A. Afirmación falsa
B. Otra afirmación falsa
C. Afirmación verdadera
D. Cuarta afirmación falsa
Correcta: C
Explicación: Aquí va la justificación.

4. Identifica el error en la siguiente situación:
A. Error correctamente identificado
B. Algo que no es un error
C. Algo que tampoco es un error
D. Confusión común
Correcta: A
"""

for line in quiz_text.split('\n'):
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

doc.add_paragraph()

# ========================= TEMA 2 (más reducido) =========================
add_styled_paragraph(doc, 'Tema 2. Segundo tema (opcional)', style='Heading 1')

doc.add_paragraph(
    'Si tu curso tiene varios temas, repite la estructura de Tema 1 con todos '
    'sus subapartados, callouts, ejemplos y quiz. La app generará un SCORM por tema. '
    'Si solo quieres un tema, borra todo este bloque "Tema 2" y los siguientes.'
)

add_styled_paragraph(doc, '2.1. Primer subapartado del segundo tema', style='Heading 2')
doc.add_paragraph('Tu contenido aquí...')

add_styled_paragraph(doc, '2.2. Segundo subapartado del segundo tema', style='Heading 2')
doc.add_paragraph('Tu contenido aquí...')

add_styled_paragraph(doc, '2.3. Quiz del segundo tema', style='Heading 2')
doc.add_paragraph('(Repite el formato de preguntas del Tema 1)')

doc.add_paragraph()

# ========================= AYUDA FINAL =========================
add_styled_paragraph(doc, '🆘 ¿Necesitas ayuda?', style='Heading 2')

doc.add_paragraph(
    'Consulta la documentación completa de la convención DOCX en la app, '
    'o ponte en contacto con el soporte. Si tu Word tiene un formato distinto al '
    'de esta plantilla, la app intentará interpretarlo igualmente: el parser es '
    'tolerante a fallos y siempre genera un SCORM válido, aunque emita avisos '
    'sobre lo que no ha podido interpretar.'
)

# ========================= GUARDAR =========================
import os
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"OK: plantilla generada en {OUTPUT}")
print(f"Tamaño: {os.path.getsize(OUTPUT)} bytes")
