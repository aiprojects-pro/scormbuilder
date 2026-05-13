"""Genera un DOCX de test con todos los casos críticos que la v0.5 debe manejar.

Incluye:
- Imagen incrustada (PNG generada en memoria)
- Texto con negrita y cursiva
- Hipervínculo dentro de una palabra ("ver vídeo" → YouTube)
- Hipervínculo a página normal
- URL suelta como texto plano
- Enlace de YouTube como texto plano
- Lista con enlaces dentro
- Tabla con enlaces y negrita
- Callout con enlace dentro
- Quiz
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

try:
    from PIL import Image, ImageDraw
    PIL_OK = True
except ImportError:
    PIL_OK = False


def _make_test_image() -> io.BytesIO:
    """Crea una imagen PNG simple en memoria."""
    if not PIL_OK:
        # PNG mínimo (1x1 rojo) hardcoded
        import base64
        data = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
        )
        return io.BytesIO(data)
    img = Image.new("RGB", (320, 180), color=(30, 80, 160))
    draw = ImageDraw.Draw(img)
    draw.text((20, 70), "IMAGEN DE PRUEBA", fill="white")
    draw.text((20, 110), "(incrustada en el Word)", fill="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def add_hyperlink(paragraph, url: str, text: str, bold: bool = False):
    """Añade un hipervínculo a un párrafo (python-docx no lo trae nativo)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Color azul tipo enlace
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main():
    doc = Document()

    # Metadatos
    doc.add_paragraph("---")
    doc.add_paragraph("TITULO: Curso de prueba v0.5")
    doc.add_paragraph("SUBTITULO: Validación de imágenes, enlaces y WCAG")
    doc.add_paragraph("AUTOR: Test automatizado")
    doc.add_paragraph("PALETA: azul")
    doc.add_paragraph("MASTERY: 70")
    doc.add_paragraph("---")

    # TEMA 1
    doc.add_heading("Tema 1. Casos críticos de extracción", level=1)

    # Subapartado 1.1 — Imágenes y formato inline
    doc.add_heading("1.1 Imágenes incrustadas y formato", level=2)
    p = doc.add_paragraph("Este párrafo tiene una palabra en ")
    p.add_run("negrita").bold = True
    p.add_run(", una en ")
    p.add_run("cursiva").italic = True
    p.add_run(" y otra ")
    r = p.add_run("combinada")
    r.bold = True; r.italic = True
    p.add_run(".")

    doc.add_paragraph("A continuación una imagen incrustada en el Word:")

    img_p = doc.add_paragraph()
    img_p.add_run().add_picture(_make_test_image(), width=Inches(3))
    doc.add_paragraph("Figura 1: imagen de prueba generada en memoria.")

    # Subapartado 1.2 — Hipervínculos
    doc.add_heading("1.2 Hipervínculos del Word", level=2)

    # Enlace dentro de una palabra
    p = doc.add_paragraph("Mira este recurso: ")
    add_hyperlink(p, "https://www.youtube.com/watch?v=dQw4w9WgXcQ", "ver vídeo de YouTube")
    p.add_run(" para entender el concepto.")

    # Enlace a web normal
    p = doc.add_paragraph("Consulta la ")
    add_hyperlink(p, "https://www.boe.es", "página del BOE")
    p.add_run(" para la normativa completa.")

    # URL suelta como texto plano
    doc.add_paragraph(
        "También puedes ir directamente a https://es.wikipedia.org/wiki/Deporte "
        "para más información."
    )

    # YouTube como texto plano (no enlace)
    doc.add_paragraph(
        "Otro vídeo recomendado (URL suelta): https://www.youtube.com/watch?v=9bZkp7q19f0"
    )

    # Subapartado 1.3 — Lista y tabla con formato
    doc.add_heading("1.3 Listas y tablas con enlaces", level=2)

    doc.add_paragraph("Recursos clave:")
    bullets = [
        ("Consulta el ", "https://moodle.org", "manual oficial de Moodle"),
        ("Estándar ", "https://www.imsglobal.org/scorm", "SCORM en IMS"),
        ("Guía ", "https://www.w3.org/TR/WCAG21/", "WCAG 2.1"),
    ]
    for prefix, url, link_text in bullets:
        p = doc.add_paragraph(prefix, style="List Bullet")
        add_hyperlink(p, url, link_text)
        p.add_run(".")

    # Tabla con enlace y negrita
    doc.add_paragraph("Tabla resumen:")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Recurso"
    table.cell(0, 1).text = "Enlace"

    cell = table.cell(1, 0)
    cell.text = ""
    r = cell.paragraphs[0].add_run("Documentación oficial")
    r.bold = True
    add_hyperlink(table.cell(1, 1).paragraphs[0], "https://docs.python.org", "docs.python.org")

    cell = table.cell(2, 0)
    cell.text = ""
    cell.paragraphs[0].add_run("WCAG").italic = True
    add_hyperlink(table.cell(2, 1).paragraphs[0], "https://www.w3.org/WAI/", "W3C WAI")

    # Subapartado 1.4 — Callouts y multimedia con sintaxis [TIPO]
    doc.add_heading("1.4 Callouts con formato", level=2)

    p = doc.add_paragraph()
    p.add_run("[CLAVE] ").bold = True
    p.add_run("Este callout tiene una palabra ")
    p.add_run("en negrita").bold = True
    p.add_run(" y un enlace dentro.")

    p = doc.add_paragraph("[ALERTA] No olvides revisar la ")
    add_hyperlink(p, "https://www.boe.es/buscar/act.php?id=BOE-A-2013-12886", "LOMCE")
    p.add_run(" antes de continuar.")

    doc.add_paragraph("[EXITO] Buena práctica documentada en el manual.")

    # Quiz
    doc.add_heading("Quiz del tema", level=2)
    doc.add_paragraph("1. ¿El parser de v0.5 extrae imágenes incrustadas en el Word?")
    doc.add_paragraph("A. No, solo si se usa el prefijo [IMAGEN]")
    doc.add_paragraph("B. Sí, automáticamente")
    doc.add_paragraph("C. Solo si pesan menos de 100KB")
    doc.add_paragraph("D. Solo en formato SVG")
    doc.add_paragraph("Correcta: B")
    doc.add_paragraph("Explicación: La extracción automática es una de las mejoras clave de v0.5.")

    doc.add_paragraph("2. ¿Qué pasa con un enlace de YouTube embebido en una palabra del Word?")
    doc.add_paragraph("A. Se pierde por completo")
    doc.add_paragraph("B. Se queda solo como texto sin enlace")
    doc.add_paragraph("C. Se transforma en un reproductor de vídeo embebido")
    doc.add_paragraph("D. Se descarga el vídeo al SCORM")
    doc.add_paragraph("Correcta: C")
    doc.add_paragraph("Explicación: La v0.5 detecta YouTube/Vimeo y embebe el reproductor.")

    out = Path("/home/claude/scormbuilder-v05/plantilla/test_v05.docx")
    doc.save(str(out))
    print(f"Generado: {out}")
    print(f"Tamaño: {out.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
