"""Parser de DOCX a estructura intermedia.

Lee un archivo Word, identifica temas, subapartados, párrafos, callouts,
ejemplos, tablas, listas y quiz, y produce una estructura JSON-compatible
que el renderer convertirá luego en HTML.

NOVEDADES v0.2:
- Detección robusta de párrafos sin estilo (style=None ya no rompe).
- Detección de temas/subapartados por patrón de texto cuando no hay
  estilos Heading 1/2 (caso muy frecuente en docx generados por terceros).
- Bloques multimedia: IMAGE, VIDEO, AUDIO, EMBED, RESOURCE.
"""
from __future__ import annotations

import re
import logging
import tempfile
import unicodedata
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Dict, Optional, Any
from enum import Enum

from docx import Document

from scorm_builder.inline import (
    ImageExtractor, ExtraBlock, process_paragraph_inline, get_plain_text, is_video_url,
)

logger = logging.getLogger(__name__)


# ============================================================
# MODELO DE ESTRUCTURA INTERMEDIA
# ============================================================

class BlockType(str, Enum):
    PARAGRAPH = "paragraph"
    HEADING_3 = "heading_3"
    HEADING_4 = "heading_4"
    LIST_BULLET = "list_bullet"
    LIST_NUMBER = "list_number"
    TABLE = "table"
    CALLOUT_KEY = "callout_key"
    CALLOUT_ALERT = "callout_alert"
    CALLOUT_SUCCESS = "callout_success"
    CALLOUT_WARN = "callout_warn"
    EXAMPLE = "example"
    QUOTE = "quote"
    DOWNLOAD = "download"
    # NUEVOS bloques multimedia
    IMAGE = "image"            # imagen JPG/PNG/SVG/WebP/GIF
    VIDEO = "video"            # vídeo MP4/WebM/Ogg
    AUDIO = "audio"            # audio MP3/Ogg/WAV
    EMBED = "embed"            # iframe (YouTube, Vimeo, H5P, etc.)
    RESOURCE = "resource"      # cualquier otro recurso descargable


@dataclass
class Block:
    """Un bloque de contenido dentro de un subapartado.

    Campos enriquecidos (v0.5):
      - `text_html`: HTML preservando negritas, cursivas, enlaces, etc. Si está
        presente, el renderer lo usa SIN volver a escapar. Cuando es None, se
        cae a `text` (plano) y se escapa.
      - `items_html`: equivalente a `text_html` para items de lista.
      - `rows_html`: equivalente para celdas de tabla.
    """
    type: BlockType
    text: str = ""
    items: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)  # solo TABLE
    extras: Dict[str, str] = field(default_factory=dict)  # ANALISIS, REFLEXION, file, src, ...
    text_html: Optional[str] = None  # v0.5: HTML inline preservado
    items_html: Optional[List[str]] = None  # v0.5: HTML inline por item
    rows_html: Optional[List[List[str]]] = None  # v0.5: HTML inline por celda


@dataclass
class Question:
    """Una pregunta del quiz.

    v0.5 Fase 2: `qtype` permite distintos tipos:
      - "multiple_choice": opción múltiple (4 opciones por defecto). DEFAULT.
      - "true_false": verdadero/falso (options = ["Verdadero", "Falso"]).
      - "fill_in": completar hueco (texto con "___", opciones = posibles palabras).
    """
    text: str
    options: List[str]
    correct_index: int  # 0=A, 1=B, 2=C, 3=D
    explanation: Optional[str] = None
    qtype: str = "multiple_choice"


@dataclass
class Subsection:
    """Un subapartado dentro de un tema (h2)."""
    id: str  # identificador URL-safe (p.ej. "l1")
    number: str  # "1.1"
    title: str
    blocks: List[Block] = field(default_factory=list)


@dataclass
class Topic:
    """Un tema completo (h1) que se convertirá en un SCORM."""
    number: int  # 1, 2, 3...
    title: str
    subsections: List[Subsection] = field(default_factory=list)
    quiz: List[Question] = field(default_factory=list)
    intro: Optional[str] = None  # texto antes del primer subapartado
    # v0.5 Fase 2: etiquetas temáticas (5-8) generadas por IA o manuales.
    # Se inyectan como <keyword> en el manifest SCORM y como chips visibles
    # bajo el título del tema en el HTML.
    tags: List[str] = field(default_factory=list)
    # v0.5 Fase 2: preguntas intercaladas por subapartado, mapeo {sub_id: [Question, ...]}.
    # Se renderizan al final de cada subapartado correspondiente.
    inline_quiz: Dict[str, List["Question"]] = field(default_factory=dict)


@dataclass
class CourseMetadata:
    """Metadatos del curso completo."""
    title: str = "Curso sin título"
    subtitle: str = ""
    author: str = ""
    sector: str = ""
    palette: str = "azul"
    mastery: int = 70
    # Sistema de puntuación ponderada (v0.3)
    weight_view: int = 40        # peso de la visualización (0-100)
    weight_quiz: int = 60        # peso del quiz (0-100)
    view_min_seconds: int = 10   # tiempo mínimo por subapartado para contar como visto
    view_strategy: str = "both"  # "scroll", "time", "both"
    # v0.5.7: colores custom de la paleta (si el usuario los personalizó al generar).
    # Si están los 3, se usan en lugar de la paleta predefinida 'palette'.
    # Si están vacíos, se usa la paleta predefinida por nombre.
    color_deep: str = ""         # cabecera (oscuro)
    color_primary: str = ""      # color primario
    color_bright: str = ""       # color brillante (acentos)


@dataclass
class CourseStructure:
    """Estructura completa de un curso parseado desde DOCX."""
    metadata: CourseMetadata
    topics: List[Topic] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    # v0.5: imágenes incrustadas extraídas del DOCX. Ruta de la carpeta y
    # lista de nombres de fichero. La carpeta debe pasarse al packager como
    # parte de recursos_dir para que las imágenes acaben dentro del SCORM.
    extracted_images_dir: Optional[str] = None
    extracted_image_files: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "metadata": asdict(self.metadata),
            "topics": [
                {
                    "number": t.number,
                    "title": t.title,
                    "intro": t.intro,
                    "tags": list(t.tags),
                    "subsections": [
                        {
                            "id": s.id,
                            "number": s.number,
                            "title": s.title,
                            "blocks": [
                                {
                                    "type": b.type.value if isinstance(b.type, BlockType) else b.type,
                                    "text": b.text,
                                    "items": b.items,
                                    "rows": b.rows,
                                    "extras": b.extras,
                                    "text_html": b.text_html,
                                    "items_html": b.items_html,
                                    "rows_html": b.rows_html,
                                }
                                for b in s.blocks
                            ],
                        }
                        for s in t.subsections
                    ],
                    "quiz": [
                        {
                            "text": q.text,
                            "options": q.options,
                            "correct_index": q.correct_index,
                            "explanation": q.explanation,
                            "qtype": q.qtype,
                        }
                        for q in t.quiz
                    ],
                    "inline_quiz": {
                        sub_id: [
                            {
                                "text": q.text,
                                "options": q.options,
                                "correct_index": q.correct_index,
                                "explanation": q.explanation,
                                "qtype": q.qtype,
                            }
                            for q in qs
                        ]
                        for sub_id, qs in t.inline_quiz.items()
                    },
                }
                for t in self.topics
            ],
            "warnings": self.warnings,
        }


# ============================================================
# PARSER
# ============================================================

# Palabras clave de bloques especiales (insensibles a tildes/mayúsculas)
CALLOUT_PREFIXES = {
    "CLAVE": BlockType.CALLOUT_KEY,
    "ALERTA": BlockType.CALLOUT_ALERT,
    "EXITO": BlockType.CALLOUT_SUCCESS,
    "ÉXITO": BlockType.CALLOUT_SUCCESS,
    "CUIDADO": BlockType.CALLOUT_WARN,
    "CITA": BlockType.QUOTE,
    "DESCARGABLE": BlockType.DOWNLOAD,
    # NUEVOS
    "IMAGEN": BlockType.IMAGE,
    "FOTO": BlockType.IMAGE,
    "VIDEO": BlockType.VIDEO,
    "VÍDEO": BlockType.VIDEO,
    "AUDIO": BlockType.AUDIO,
    "EMBED": BlockType.EMBED,
    "INCRUSTAR": BlockType.EMBED,
    "RECURSO": BlockType.RESOURCE,
}

QUIZ_HEADING_KEYWORDS = ["quiz", "test", "evaluación", "evaluacion", "comprueba"]
METADATA_KEYS = ["TITULO", "TÍTULO", "SUBTITULO", "SUBTÍTULO", "AUTOR", "SECTOR", "PALETA", "MASTERY"]

# Patrones de respaldo para detectar Heading 1 / Heading 2 sin estilos
HEADING1_PATTERNS = [
    re.compile(r"^\s*tema\s+\d+[\.\s\-:]", re.IGNORECASE),
    re.compile(r"^\s*módulo\s+\d+[\.\s\-:]", re.IGNORECASE),
    re.compile(r"^\s*modulo\s+\d+[\.\s\-:]", re.IGNORECASE),
    re.compile(r"^\s*unidad\s+\d+[\.\s\-:]", re.IGNORECASE),
    re.compile(r"^\s*capítulo\s+\d+[\.\s\-:]", re.IGNORECASE),
    re.compile(r"^\s*capitulo\s+\d+[\.\s\-:]", re.IGNORECASE),
    re.compile(r"^\s*lección\s+\d+[\.\s\-:]", re.IGNORECASE),
    re.compile(r"^\s*leccion\s+\d+[\.\s\-:]", re.IGNORECASE),
]

HEADING2_PATTERN = re.compile(r"^\s*(\d+)\.(\d+)\.?\s+\S")


def _slugify(text: str) -> str:
    """Convierte un texto en un id URL-safe."""
    text = unicodedata.normalize("NFKD", text)
    text = text.encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
    return text or "section"


def _safe_style_name(p) -> str:
    """Devuelve el nombre del estilo de un párrafo, o cadena vacía si no hay.

    python-docx puede devolver None en p.style si el párrafo se ha creado sin
    estilo asignado o el estilo se ha eliminado. Antes esto rompía el parser.
    """
    try:
        if p is None:
            return ""
        st = p.style
        if st is None:
            return ""
        name = st.name
        return name or ""
    except Exception:
        return ""


def _detect_callout_prefix(text: str) -> Optional[BlockType]:
    """Detecta si un párrafo empieza con [TIPO] y devuelve el BlockType."""
    match = re.match(r"^\s*\[(\w+(?:\s\w+)*)\]\s*", text)
    if not match:
        return None
    tag = match.group(1).upper().strip()
    return CALLOUT_PREFIXES.get(tag)


def _strip_callout_prefix(text: str) -> str:
    """Elimina el prefijo [TIPO] del texto."""
    return re.sub(r"^\s*\[\w+(?:\s\w+)*\]\s*", "", text, count=1)


def _strip_callout_prefix_html(html: str) -> str:
    """Elimina el prefijo [TIPO] del HTML inline, incluso si está dentro de
    tags de formato (`<strong>[CLAVE]</strong>`). Limpia las tags vacías que
    queden tras el borrado."""
    # 1) Eliminar el [TIPO] esté donde esté (suele venir al principio).
    cleaned = re.sub(
        r"\s*\[\w+(?:\s\w+)*\]\s*",
        "",
        html, count=1,
    )
    # 2) Limpiar tags de formato vacías al principio: <strong></strong>,
    #    <em></em>, <u></u>, <s></s>. Repetir por si están anidadas.
    for _ in range(3):
        new = re.sub(
            r"^\s*<(strong|em|u|s)>\s*</\1>\s*",
            "",
            cleaned,
        )
        if new == cleaned:
            break
        cleaned = new
    return cleaned.strip()


def _looks_like_heading1(text: str) -> bool:
    """Devuelve True si el texto encaja con un patrón típico de Heading 1."""
    if not text:
        return False
    return any(p.match(text) for p in HEADING1_PATTERNS)


def _looks_like_heading2(text: str) -> bool:
    """Devuelve True si el texto encaja con un patrón típico de Heading 2 (N.M)."""
    if not text or len(text) > 200:
        return False
    return bool(HEADING2_PATTERN.match(text))


def _detect_metadata(paragraphs: List[Any], course: CourseStructure) -> int:
    """Si encuentra un bloque de metadatos al principio (entre --- y ---),
    los aplica al course.metadata y devuelve el índice del párrafo donde
    termina ese bloque (para saltarlo). Si no hay metadatos, devuelve 0."""
    if not paragraphs:
        return 0

    # Buscar primer "---"
    start = None
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if text == "---":
            start = i
            break
        # Si encontramos un Heading 1 (por estilo o patrón) antes, no hay metadatos
        style_name = _safe_style_name(p)
        if style_name.startswith("Heading 1") or _looks_like_heading1(text):
            return 0
        # Si encontramos texto significativo antes del ---, no hay metadatos
        if text and not text.startswith("---"):
            return 0

    if start is None:
        return 0

    # Buscar segundo "---"
    end = None
    for i in range(start + 1, len(paragraphs)):
        if paragraphs[i].text.strip() == "---":
            end = i
            break

    if end is None:
        return 0

    # Procesar líneas de metadatos
    for j in range(start + 1, end):
        line = paragraphs[j].text.strip()
        if ":" not in line:
            continue
        key, _, value = line.partition(":")
        key = key.strip().upper()
        value = value.strip()
        if key in ("TITULO", "TÍTULO"):
            course.metadata.title = value
        elif key in ("SUBTITULO", "SUBTÍTULO"):
            course.metadata.subtitle = value
        elif key == "AUTOR":
            course.metadata.author = value
        elif key == "SECTOR":
            course.metadata.sector = value
        elif key == "PALETA":
            course.metadata.palette = value.lower()
        elif key == "MASTERY":
            try:
                course.metadata.mastery = int(value)
            except ValueError:
                pass
        elif key in ("PESO_VISUALIZACION", "PESO_VISUALIZACIÓN", "WEIGHT_VIEW"):
            try:
                course.metadata.weight_view = max(0, min(100, int(value)))
            except ValueError:
                pass
        elif key in ("PESO_QUIZ", "WEIGHT_QUIZ"):
            try:
                course.metadata.weight_quiz = max(0, min(100, int(value)))
            except ValueError:
                pass
        elif key in ("TIEMPO_MINIMO", "TIEMPO_MÍNIMO", "VIEW_MIN_SECONDS"):
            try:
                course.metadata.view_min_seconds = max(0, int(value))
            except ValueError:
                pass
        elif key in ("ESTRATEGIA_VISTA", "VIEW_STRATEGY"):
            v = value.lower().strip()
            if v in ("scroll", "time", "tiempo", "both", "ambos", "mixto"):
                course.metadata.view_strategy = {
                    "tiempo": "time", "ambos": "both", "mixto": "both",
                }.get(v, v)

    return end + 1


def _is_quiz_heading(text: str) -> bool:
    """Detecta si un Heading 2 es un encabezado de quiz."""
    text_lower = text.lower()
    return any(kw in text_lower for kw in QUIZ_HEADING_KEYWORDS)


def _parse_quiz_block(quiz_text: str) -> List[Question]:
    """Parsea el texto crudo de un quiz y extrae las preguntas."""
    questions: List[Question] = []
    # Dividir por preguntas: cada pregunta empieza con "N." al inicio de línea
    chunks = re.split(r"\n(?=\d+\.\s)", quiz_text.strip())

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        lines = [l.strip() for l in chunk.split("\n") if l.strip()]
        if not lines:
            continue

        # Primera línea: pregunta
        first = lines[0]
        m = re.match(r"^\d+\.\s*(.+)", first)
        if not m:
            continue
        question_text = m.group(1).strip()

        options: List[str] = []
        correct_index = -1
        explanation: Optional[str] = None

        for line in lines[1:]:
            mo = re.match(r"^([A-Da-d])[\.\)]\s*(.+)", line)
            if mo:
                options.append(mo.group(2).strip())
                continue
            mc = re.match(r"^Correcta\s*:\s*([A-Da-d])\s*$", line, re.IGNORECASE)
            if mc:
                letter = mc.group(1).upper()
                correct_index = ord(letter) - ord("A")
                continue
            me = re.match(r"^Explicaci[óo]n\s*:\s*(.+)", line, re.IGNORECASE)
            if me:
                explanation = me.group(1).strip()
                continue
            # Líneas adicionales tras la explicación: las concatenamos
            if explanation:
                explanation += " " + line

        if len(options) >= 2 and 0 <= correct_index < len(options):
            questions.append(Question(
                text=question_text,
                options=options,
                correct_index=correct_index,
                explanation=explanation,
            ))

    return questions


def _extract_table_rows(table) -> List[List[str]]:
    """Extrae las filas de una tabla DOCX como listas de strings."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return rows


def _extract_table_rows_html(table, doc, extractor) -> tuple[List[List[str]], List[ExtraBlock]]:
    """Extrae las filas como HTML enriquecido (cada celda preserva enlaces y formato).
    Devuelve (rows_html, extras_no_inline). Los extras (imágenes en celdas) se
    pierden en este formato porque ir a la celda complica mucho; los devolvemos
    para emitirlos tras la tabla.
    """
    rows_html: List[List[str]] = []
    extras_collected: List[ExtraBlock] = []
    for row in table.rows:
        row_html = []
        for cell in row.cells:
            cell_parts = []
            for p in cell.paragraphs:
                inline_html, extras = process_paragraph_inline(p, doc, extractor)
                if inline_html.strip():
                    cell_parts.append(inline_html.strip())
                extras_collected.extend(extras)
            row_html.append("<br>".join(cell_parts))
        rows_html.append(row_html)
    return rows_html, extras_collected


def _is_list_paragraph(p) -> Optional[BlockType]:
    """Detecta si un párrafo es de lista y de qué tipo."""
    style_name = _safe_style_name(p)
    if style_name.startswith("List Bullet"):
        return BlockType.LIST_BULLET
    if style_name.startswith("List Number"):
        return BlockType.LIST_NUMBER
    # Detección por numId del XML (más fiable)
    try:
        pPr = p._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        if pPr is not None:
            numPr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
            if numPr is not None:
                return BlockType.LIST_BULLET  # asumimos bullets si no podemos saber el tipo
    except Exception:
        pass
    return None


def _build_media_block(callout_type: BlockType, clean_text: str) -> Optional[Block]:
    """Construye un Block multimedia a partir del texto tras el prefijo.

    Sintaxis aceptada:
        [IMAGEN] Pie de imagen | nombre_archivo.png
        [VIDEO]  Título | https://youtu.be/abc        (URL externa)
        [VIDEO]  Título | mi_video.mp4                (archivo subido)
        [AUDIO]  Locución | locucion.mp3
        [EMBED]  Cuestionario H5P | https://h5p.org/embed/...
        [RECURSO] Plantilla Excel | plantilla.xlsx
    """
    if "|" in clean_text:
        label, _, src = clean_text.partition("|")
        label = label.strip()
        src = src.strip()
    else:
        label = ""
        src = clean_text.strip()
    if not src:
        return None
    return Block(
        type=callout_type,
        text=label,
        extras={"src": src, "file": src},
    )


def parse_docx(
    path: str | Path,
    default_palette: str = "azul",
    images_dir: str | Path | None = None,
) -> CourseStructure:
    """Parsea un archivo DOCX y devuelve una CourseStructure.

    Args:
        path: ruta al DOCX
        default_palette: paleta por defecto si no se indica en metadatos
        images_dir: carpeta donde extraer las imágenes incrustadas del DOCX.
            Si es None, se crea una temporal. La ruta queda en
            `course.extracted_images_dir` para que el caller pueda incluirla
            como recursos en el SCORM.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {path}")

    doc = Document(str(path))

    course = CourseStructure(metadata=CourseMetadata(palette=default_palette))

    # v0.5: extractor de imágenes incrustadas
    if images_dir is None:
        images_dir = Path(tempfile.mkdtemp(prefix="scormbuilder_imgs_"))
    else:
        images_dir = Path(images_dir)
        images_dir.mkdir(parents=True, exist_ok=True)
    extractor = ImageExtractor(doc, images_dir)
    course.extracted_images_dir = str(images_dir)

    # Iterar todos los bloques del cuerpo (párrafos + tablas en orden)
    body_elements = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            for p in doc.paragraphs:
                if p._p == child:
                    body_elements.append(("p", p))
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._tbl == child:
                    body_elements.append(("tbl", t))
                    break

    paragraphs = [el[1] for el in body_elements if el[0] == "p"]

    # Procesar metadatos al inicio (si los hay)
    skip_until = _detect_metadata(paragraphs, course)

    # Estado del parser
    current_topic: Optional[Topic] = None
    current_subsection: Optional[Subsection] = None
    in_quiz = False
    quiz_buffer: List[str] = []
    list_buffer: List[str] = []  # texto plano (compatibilidad)
    list_buffer_html: List[str] = []  # html enriquecido por item
    list_buffer_extras: List[ExtraBlock] = []  # imágenes / vídeos pendientes
    list_type: Optional[BlockType] = None
    intro_buffer: List[str] = []  # texto antes del primer h2 de cada tema

    p_index = 0

    def flush_list():
        nonlocal list_buffer, list_buffer_html, list_buffer_extras, list_type
        if list_buffer and current_subsection and list_type:
            block = Block(type=list_type, items=list_buffer.copy())
            # Si hay html enriquecido en algún item, lo añadimos
            if any(h is not None for h in list_buffer_html):
                # Para items sin html (raros) caemos a escape del texto plano
                import html as _html
                block.items_html = [
                    h if h is not None else _html.escape(t, quote=False)
                    for h, t in zip(list_buffer_html, list_buffer)
                ]
            current_subsection.blocks.append(block)
            # Emitir extras pendientes (imágenes/vídeos que iban en los items)
            for ex in list_buffer_extras:
                _emit_extra_block(current_subsection, ex)
            list_buffer = []
            list_buffer_html = []
            list_buffer_extras = []
            list_type = None

    def _emit_extra_block(target_sub: "Subsection", ex: ExtraBlock):
        """Convierte un ExtraBlock (imagen/vídeo extraído) en un Block y lo añade."""
        if ex.type == "image":
            target_sub.blocks.append(Block(
                type=BlockType.IMAGE,
                text=ex.caption,
                extras={"src": ex.src, "file": ex.src,
                        **{k: str(v) for k, v in ex.extras.items()}},
            ))
        elif ex.type == "video_embed":
            target_sub.blocks.append(Block(
                type=BlockType.VIDEO,
                text=ex.caption,
                extras={"src": ex.src, "file": ex.src},
            ))

    def flush_quiz():
        nonlocal quiz_buffer, in_quiz
        if quiz_buffer and current_topic:
            text = "\n".join(quiz_buffer)
            questions = _parse_quiz_block(text)
            current_topic.quiz.extend(questions)
            if not questions:
                course.warnings.append(f"El quiz del tema {current_topic.number} no contiene preguntas válidas.")
            quiz_buffer = []
        in_quiz = False

    def flush_intro():
        nonlocal intro_buffer
        if intro_buffer and current_topic:
            current_topic.intro = " ".join(intro_buffer).strip()
            intro_buffer = []

    for kind, el in body_elements:
        # Saltar bloque de metadatos inicial
        if kind == "p":
            try:
                p_index = paragraphs.index(el)
            except ValueError:
                pass
            if p_index < skip_until:
                continue

        if kind == "tbl":
            # Tabla
            if current_subsection:
                flush_list()
                rows = _extract_table_rows(el)
                if rows:
                    rows_html, table_extras = _extract_table_rows_html(el, doc, extractor)
                    current_subsection.blocks.append(
                        Block(type=BlockType.TABLE, rows=rows, rows_html=rows_html)
                    )
                    # Si había imágenes dentro de la tabla, emitirlas debajo
                    for ex in table_extras:
                        _emit_extra_block(current_subsection, ex)
            continue

        # Es un párrafo
        p = el
        # v0.5: get_plain_text es como p.text pero más consistente con
        # cómo trataremos hyperlinks. Para detectar headings, callouts, etc.
        text = get_plain_text(p).strip()
        style_name = _safe_style_name(p)

        # ---- HEADINGS (con respaldo por patrón de texto) ----
        is_h1 = style_name.startswith("Heading 1") or _looks_like_heading1(text)
        # Si el texto coincide con un patrón H1, prevalece sobre H2 aunque tenga
        # estilo Heading 2 mal puesto (caso de docx extraídos de PDF).
        # Detección de h2: estilo Heading 2, patrón N.M., o ser un encabezado de quiz
        # corto como "Quiz", "Test", "Evaluación" en línea suelta.
        is_quiz_kw_line = (
            not is_h1 and 1 <= len(text.split()) <= 4 and _is_quiz_heading(text)
        )
        is_h2 = (
            (style_name.startswith("Heading 2") or _looks_like_heading2(text) or is_quiz_kw_line)
            and not is_h1
        )
        is_h3 = style_name.startswith("Heading 3") and not is_h1 and not is_h2
        is_h4 = style_name.startswith("Heading 4") and not is_h1 and not is_h2 and not is_h3

        # Detectar Heading 1 → nuevo tema
        if is_h1:
            flush_list()
            flush_quiz()
            flush_intro()
            topic_number = len(course.topics) + 1
            # Limpiar título: quitar "Tema N.", "Módulo N.", etc.
            title_clean = re.sub(
                r"^\s*(tema|módulo|modulo|unidad|capítulo|capitulo|lección|leccion)\s+\d+[\.\s\-:]\s*",
                "", text, flags=re.IGNORECASE,
            ).strip()
            current_topic = Topic(
                number=topic_number,
                title=title_clean if title_clean else f"Tema {topic_number}",
            )
            course.topics.append(current_topic)
            current_subsection = None
            in_quiz = False
            continue

        # Si todavía no hay tema, ignorar
        if current_topic is None:
            if text:
                # Si hay contenido fuera de un tema, lo descartamos con warning
                course.warnings.append(f"Texto fuera de un tema descartado: '{text[:50]}...'")
            continue

        # Detectar Heading 2 → subapartado o quiz
        if is_h2:
            flush_list()
            flush_quiz()
            flush_intro()
            if _is_quiz_heading(text):
                in_quiz = True
                continue
            # Subapartado normal
            sub_number = f"{current_topic.number}.{len(current_topic.subsections) + 1}"
            # Limpiar el número del título si lo lleva
            title_clean = re.sub(r"^\s*\d+\.\d+\.?\s*", "", text).strip()
            sub_id = f"l{len(current_topic.subsections) + 1}"
            current_subsection = Subsection(
                id=sub_id,
                number=sub_number,
                title=title_clean,
            )
            current_topic.subsections.append(current_subsection)
            continue

        # Si estamos dentro del quiz, acumular
        if in_quiz:
            if text:
                quiz_buffer.append(text)
            continue

        # Heading 3 o 4
        if is_h3:
            flush_list()
            if current_subsection:
                current_subsection.blocks.append(Block(type=BlockType.HEADING_3, text=text))
            continue
        if is_h4:
            flush_list()
            if current_subsection:
                current_subsection.blocks.append(Block(type=BlockType.HEADING_4, text=text))
            continue

        # v0.5: Antes de cualquier otra cosa, procesar inline el párrafo.
        # Esto extrae imágenes, hipervínculos y URLs sueltas. Si el párrafo SOLO
        # contiene una imagen incrustada (sin texto), saldrá text="" pero un
        # ExtraBlock de imagen.
        inline_html, extras = process_paragraph_inline(p, doc, extractor)
        # Detectar si el párrafo está vacío de texto pero trae extras
        text_clean = re.sub(r"<[^>]+>", "", inline_html).strip()

        # Lista (bullet o number)
        list_kind = _is_list_paragraph(p)
        if list_kind and text:
            if list_type is not None and list_type != list_kind:
                flush_list()
            list_type = list_kind
            list_buffer.append(text)
            list_buffer_html.append(inline_html if text_clean else None)
            # Los extras del item se acumulan y se emiten al cerrar la lista
            list_buffer_extras.extend(extras)
            continue
        else:
            flush_list()

        # Si no hay subapartado todavía, añadir al intro del tema
        if current_subsection is None:
            if text and not text.startswith("---"):
                intro_buffer.append(text)
            continue

        # Caso: párrafo sin texto pero con extras (imagen suelta en el Word)
        if not text and extras:
            for ex in extras:
                _emit_extra_block(current_subsection, ex)
            continue

        # Detectar callout / multimedia
        if text:
            callout_type = _detect_callout_prefix(text)
            if callout_type:
                clean = _strip_callout_prefix(text)
                # Para el callout también queremos preservar el HTML inline,
                # pero el prefijo [TIPO] está en el HTML — lo quitamos del HTML.
                clean_html = _strip_callout_prefix_html(inline_html)
                # Casos especiales con archivo asociado
                if callout_type == BlockType.QUOTE:
                    current_subsection.blocks.append(
                        Block(type=callout_type, text=clean, text_html=clean_html)
                    )
                elif callout_type in (
                    BlockType.DOWNLOAD, BlockType.IMAGE, BlockType.VIDEO,
                    BlockType.AUDIO, BlockType.EMBED, BlockType.RESOURCE,
                ):
                    media_block = _build_media_block(callout_type, clean)
                    if media_block:
                        current_subsection.blocks.append(media_block)
                    # También emitimos los extras del propio párrafo
                    for ex in extras:
                        _emit_extra_block(current_subsection, ex)
                    continue
                else:
                    current_subsection.blocks.append(
                        Block(type=callout_type, text=clean, text_html=clean_html)
                    )
                # Emitir extras (vídeos YouTube, imágenes) detectados en el callout
                for ex in extras:
                    _emit_extra_block(current_subsection, ex)
                continue

            # Párrafo normal con HTML enriquecido
            current_subsection.blocks.append(
                Block(type=BlockType.PARAGRAPH, text=text, text_html=inline_html)
            )
            # Emitir extras (imágenes/vídeos) que iban junto al texto
            for ex in extras:
                _emit_extra_block(current_subsection, ex)

    # Flush final
    flush_list()
    flush_quiz()
    flush_intro()

    # Validaciones finales
    if not course.topics:
        course.warnings.append("No se ha detectado ningún tema (Heading 1) en el documento.")

    for t in course.topics:
        if not t.subsections:
            course.warnings.append(f"El tema '{t.title}' no tiene subapartados (Heading 2).")
        if not t.quiz:
            course.warnings.append(f"El tema '{t.title}' no tiene quiz definido. Se generará SCORM sin evaluación o con preguntas por IA.")

    # Normalizar pesos del sistema de puntuación
    _normalize_weights(course)

    # v0.5: registrar los archivos de imagen extraídos del DOCX
    course.extracted_image_files = extractor.extracted_filenames()

    return course


def _normalize_weights(course: CourseStructure) -> None:
    """Normaliza los pesos de visualización y quiz a 100% en total.

    Reglas:
    - Si la suma no es 100, se reescala proporcionalmente.
    - Si la suma es 0 (caso absurdo), se aplica el default 40/60.
    - Si algún tema NO tiene quiz, ese tema usará 100% visualización en runtime;
      se añade un warning explícito por tema.
    """
    md = course.metadata
    total = md.weight_view + md.weight_quiz
    if total <= 0:
        md.weight_view = 40
        md.weight_quiz = 60
    elif total != 100:
        # Reescalar manteniendo la proporción
        scale = 100.0 / total
        md.weight_view = round(md.weight_view * scale)
        md.weight_quiz = 100 - md.weight_view  # garantizar suma exacta = 100
        course.warnings.append(
            f"Los pesos del sistema de puntuación no sumaban 100; se han reescalado a "
            f"{md.weight_view}% visualización + {md.weight_quiz}% quiz."
        )

    # Aviso por tema sin quiz: redistribución automática
    topics_without_quiz = [t for t in course.topics if not t.quiz]
    if topics_without_quiz and md.weight_quiz > 0:
        names = ", ".join(f"'{t.title}'" for t in topics_without_quiz[:3])
        more = f" y {len(topics_without_quiz) - 3} más" if len(topics_without_quiz) > 3 else ""
        course.warnings.append(
            f"Los temas {names}{more} no tienen quiz: la nota se calculará "
            f"solo por visualización (100%) en esos temas."
        )
