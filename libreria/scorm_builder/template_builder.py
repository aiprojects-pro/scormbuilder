"""Generador de plantilla Word "moderna" (v0.5 Fase 5).

A diferencia del script suelto `plantilla/generar_plantilla.py`, este módulo
forma parte de la librería y se puede invocar desde la app local o desde la
CLI para que el cliente descargue una plantilla rellenable.

Diferencias respecto a la versión anterior:
- Sintaxis moderna con bloques `[CLAVE]`, `[ALERTA]`, etc. claramente
  diferenciados por colores de fondo.
- Sección de **metadatos prellenada** entre `---` con todos los campos.
- Ejemplo de **hipervínculo dentro de palabra** ("ver vídeo" enlazado).
- Ejemplo de **imagen incrustada de muestra** (placeholder coloreado).
- Sección de quiz con 3 preguntas de ejemplo de los tres tipos
  (multiple_choice, true_false, fill_in) — aunque el parser de v0.5 solo
  detecta multiple_choice del texto, así el usuario aprende los formatos.
- Sección final de **guía rápida** explicando cada elemento.

Uso:
    from scorm_builder.template_builder import build_modern_template
    build_modern_template("plantilla.docx")
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# UTILIDADES DE FORMATO
# ============================================================

def _set_paragraph_bg(paragraph, hex_color: str) -> None:
    """Aplica color de fondo a un párrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_callout(doc, prefix: str, text: str, bg: str, text_color: str = "1F2937") -> None:
    """Añade un párrafo callout con fondo de color."""
    p = doc.add_paragraph()
    _set_paragraph_bg(p, bg)
    run_prefix = p.add_run(f"[{prefix}] ")
    run_prefix.bold = True
    run_prefix.font.color.rgb = RGBColor.from_string(text_color)
    run_text = p.add_run(text)
    run_text.font.color.rgb = RGBColor.from_string(text_color)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Inserta un hipervínculo dentro de un párrafo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _make_placeholder_image() -> io.BytesIO:
    """Genera un PNG placeholder simple (sin Pillow para ser robusto)."""
    try:
        from PIL import Image, ImageDraw, ImageFont  # type: ignore
        img = Image.new("RGB", (480, 270), color=(37, 99, 235))
        draw = ImageDraw.Draw(img)
        draw.rectangle([0, 0, 479, 269], outline=(255, 255, 255), width=4)
        try:
            font = ImageFont.truetype("DejaVuSans-Bold.ttf", 28)
        except Exception:
            font = ImageFont.load_default()
        text = "IMAGEN DE\nEJEMPLO"
        draw.multiline_text((140, 90), text, fill=(255, 255, 255),
                             font=font, align="center", spacing=10)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except ImportError:
        # PNG 1×1 azul como fallback mínimo
        import base64
        png_b64 = (
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAA"
            "DUlEQVR42mNkYPhfDwAChwGAuyKW2gAAAABJRU5ErkJggg=="
        )
        return io.BytesIO(base64.b64decode(png_b64))


# ============================================================
# GENERADOR PRINCIPAL
# ============================================================

def build_modern_template(
    output_path: str | Path,
    *,
    course_title: str = "Mi curso",
    author: str = "Tu nombre",
    sector: str = "",
    palette: str = "azul",
) -> Path:
    """Genera una plantilla DOCX rellenable y devuelve la ruta del archivo."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Márgenes razonables
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ----------------------------------------------
    # PORTADA DE INSTRUCCIONES
    # ----------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PLANTILLA DE CURSO SCORM")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0A2540")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sustituye el contenido de ejemplo por el tuyo. Conserva los estilos y palabras clave entre [CORCHETES].")
    rs.italic = True
    rs.font.color.rgb = RGBColor.from_string("64748B")
    rs.font.size = Pt(11)

    doc.add_paragraph()  # espaciador
    p = doc.add_paragraph()
    p.add_run("📌 Antes de empezar:").bold = True
    doc.add_paragraph("1. No borres la sección de metadatos entre las dos líneas '---'. Edita los valores.")
    doc.add_paragraph("2. Cada Tema empieza con un Título 1 (\"Tema N. Título\"). Cada subapartado con un Título 2 numerado (\"1.1. Tema...\").")
    doc.add_paragraph("3. Mantén las palabras clave entre corchetes ([CLAVE], [ALERTA], [EXITO], [CUIDADO], [CITA]).")
    doc.add_paragraph("4. Puedes pegar imágenes directamente con Ctrl+V o Cmd+V. Se extraerán al SCORM automáticamente.")
    doc.add_paragraph("5. Los enlaces de YouTube se convierten en reproductor automáticamente.")
    doc.add_paragraph("6. Al final hay una guía resumen con todos los formatos disponibles.")

    doc.add_page_break()

    # ----------------------------------------------
    # METADATOS
    # ----------------------------------------------
    doc.add_paragraph("---")
    doc.add_paragraph(f"TITULO: {course_title}")
    doc.add_paragraph("SUBTITULO: Subtítulo opcional del curso")
    doc.add_paragraph(f"AUTOR: {author}")
    doc.add_paragraph(f"SECTOR: {sector or 'sector / ámbito del curso'}")
    doc.add_paragraph(f"PALETA: {palette}")
    doc.add_paragraph("MASTERY: 70")
    doc.add_paragraph("PESO_VISUALIZACION: 40")
    doc.add_paragraph("PESO_QUIZ: 60")
    doc.add_paragraph("TIEMPO_MINIMO: 10")
    doc.add_paragraph("ESTRATEGIA_VISTA: both")
    doc.add_paragraph("---")

    # ----------------------------------------------
    # TEMA 1: PRIMER TEMA DE EJEMPLO
    # ----------------------------------------------
    doc.add_heading("Tema 1. Sustituye este título por el de tu primer tema", level=1)

    doc.add_paragraph(
        "Aquí escribes el primer párrafo de contenido. Puedes redactar tantos párrafos "
        "como necesites. Cada párrafo se renderiza como un bloque de texto separado en "
        "el SCORM."
    )

    # Subapartado 1.1
    doc.add_heading("1.1. Primer subapartado del tema", level=2)

    p = doc.add_paragraph(
        "Las palabras importantes puedes ponerlas en "
    )
    p.add_run("negrita").bold = True
    p.add_run(" o en ")
    p.add_run("cursiva").italic = True
    p.add_run(" y se preservarán en el SCORM. Los hipervínculos también: por ejemplo, ")
    _add_hyperlink(p, "https://www.w3.org/TR/WCAG21/", "consulta la WCAG 2.1")
    p.add_run(" para más información de accesibilidad.")

    # Ejemplos de los 5 callouts (con colores)
    _add_callout(doc, "CLAVE",
        "Aquí va una idea clave que el alumno debe retener. Se renderiza con fondo "
        "azul claro y un icono de información. Úsalo para definiciones y conceptos centrales.",
        bg="DBEAFE", text_color="1E3A8A")

    _add_callout(doc, "ALERTA",
        "Aviso importante o riesgo a evitar. Fondo rojo claro, icono de exclamación. "
        "Para prohibiciones, errores graves y situaciones que pueden tener consecuencias serias.",
        bg="FEE2E2", text_color="7F1D1D")

    _add_callout(doc, "EXITO",
        "Buena práctica o caso correcto. Fondo verde claro, icono de check. "
        "Para reforzar comportamientos deseables y casos modelo.",
        bg="D1FAE5", text_color="064E3B")

    _add_callout(doc, "CUIDADO",
        "Precaución suave. Fondo amarillo claro, icono de aviso. Para advertencias "
        "moderadas y aspectos a tener en cuenta sin dramatismo.",
        bg="FEF3C7", text_color="78350F")

    # Subapartado 1.2: lista, tabla, imagen, vídeo
    doc.add_heading("1.2. Segundo subapartado con multimedia", level=2)

    doc.add_paragraph("Las listas con viñetas se mantienen tal cual:")
    for item in ["Primer ítem de la lista",
                 "Segundo ítem con una palabra en negrita destacada",
                 "Tercer ítem de la lista"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Las listas numeradas también:")
    for item in ["Paso uno", "Paso dos", "Paso tres"]:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph("Imagen incrustada de ejemplo (puedes pegar cualquier imagen con Ctrl+V):")
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(_make_placeholder_image(), width=Inches(4))

    fig = doc.add_paragraph()
    fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_run = fig.add_run("Figura: imagen de ejemplo (reemplázala por la tuya).")
    fig_run.italic = True
    fig_run.font.color.rgb = RGBColor.from_string("64748B")
    fig_run.font.size = Pt(10)

    # Vídeo embebido (URL suelta)
    doc.add_paragraph(
        "Para insertar un vídeo de YouTube, pega la URL como texto plano o crea un "
        "hipervínculo. Ejemplo (URL plana, se convierte automáticamente en reproductor):"
    )
    doc.add_paragraph("https://www.youtube.com/watch?v=dQw4w9WgXcQ")

    # Tabla
    doc.add_paragraph("Tabla de ejemplo (la primera fila se trata como cabecera):")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Concepto", "Descripción", "Ejemplo"]
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    table.cell(1, 0).text = "Bloque clave"
    table.cell(1, 1).text = "Idea fundamental a retener"
    table.cell(1, 2).text = "[CLAVE] La accesibilidad..."
    table.cell(2, 0).text = "Cita textual"
    table.cell(2, 1).text = "Referencia a una norma"
    table.cell(2, 2).text = "[CITA] FUENTE: BOE..."

    # Cita
    _add_callout(doc, "CITA",
        'FUENTE: Real Decreto 1112/2018, artículo 1\n"Las webs y aplicaciones para móviles del sector público deberán cumplir con los requisitos de accesibilidad..."',
        bg="EEF2FF", text_color="312E81")

    # Quiz del tema 1
    doc.add_heading("Quiz del tema", level=2)

    doc.add_paragraph(
        "Escribe aquí entre 4 y 10 preguntas tipo test. Cada una con 4 opciones (A-D) "
        "y la indicación 'Correcta: X'. La explicación es opcional pero recomendada."
    )

    doc.add_paragraph("1. ¿Cuántas opciones debe tener cada pregunta del quiz?")
    doc.add_paragraph("A. Dos opciones")
    doc.add_paragraph("B. Cuatro opciones")
    doc.add_paragraph("C. Seis opciones")
    doc.add_paragraph("D. Cualquier número")
    doc.add_paragraph("Correcta: B")
    doc.add_paragraph("Explicación: La convención exige cuatro opciones por pregunta.")

    doc.add_paragraph()  # separador

    doc.add_paragraph("2. ¿Qué prefijo se usa para señalar una buena práctica?")
    doc.add_paragraph("A. [CLAVE]")
    doc.add_paragraph("B. [ALERTA]")
    doc.add_paragraph("C. [EXITO]")
    doc.add_paragraph("D. [CUIDADO]")
    doc.add_paragraph("Correcta: C")
    doc.add_paragraph("Explicación: [EXITO] señala buenas prácticas y se renderiza con fondo verde.")

    # ----------------------------------------------
    # TEMA 2 (opcional, sirve de ejemplo de cómo se separan)
    # ----------------------------------------------
    doc.add_page_break()
    doc.add_heading("Tema 2. Segundo tema (opcional, borra si solo quieres uno)", level=1)
    doc.add_heading("2.1. Subapartado del segundo tema", level=2)
    doc.add_paragraph("Cada nuevo Título 1 inicia un tema nuevo y genera un SCORM independiente.")
    doc.add_heading("Quiz del tema", level=2)
    doc.add_paragraph("1. ¿Cuántos SCORMs se generan si tu Word tiene 3 Títulos 1?")
    doc.add_paragraph("A. Uno")
    doc.add_paragraph("B. Dos")
    doc.add_paragraph("C. Tres")
    doc.add_paragraph("D. Depende del tamaño")
    doc.add_paragraph("Correcta: C")
    doc.add_paragraph("Explicación: Cada Título 1 = un tema = un SCORM independiente.")

    # ----------------------------------------------
    # GUÍA RÁPIDA AL FINAL
    # ----------------------------------------------
    doc.add_page_break()
    h = doc.add_paragraph()
    hr = h.add_run("📖 GUÍA RÁPIDA DE FORMATOS")
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = RGBColor.from_string("0A2540")

    doc.add_paragraph()  # espacio
    doc.add_paragraph().add_run("Estructura").bold = True
    doc.add_paragraph("• Título 1 = nuevo tema (genera un SCORM)")
    doc.add_paragraph("• Título 2 = subapartado dentro del tema (1.1, 1.2…)")
    doc.add_paragraph("• Título 3 / 4 = subsecciones menores")
    doc.add_paragraph("• Título 2 con 'Quiz' o 'Test' en el nombre = bloque de evaluación")

    doc.add_paragraph().add_run("Bloques destacados").bold = True
    doc.add_paragraph("• [CLAVE] Idea importante a retener (azul)")
    doc.add_paragraph("• [ALERTA] Aviso grave o prohibición (rojo)")
    doc.add_paragraph("• [EXITO] Buena práctica (verde)")
    doc.add_paragraph("• [CUIDADO] Precaución moderada (amarillo)")
    doc.add_paragraph("• [CITA] Cita textual de norma o autor (con FUENTE:)")

    doc.add_paragraph().add_run("Multimedia").bold = True
    doc.add_paragraph("• Imágenes: pega directamente en el Word, se extraen automáticamente")
    doc.add_paragraph("• Vídeos YouTube/Vimeo: pega la URL como texto o hipervínculo")
    doc.add_paragraph("• Hipervínculos: se preservan tal cual en el SCORM")
    doc.add_paragraph("• [DESCARGABLE] Nombre del recurso | archivo.pdf (lo subes aparte)")

    doc.add_paragraph().add_run("Formato del quiz").bold = True
    doc.add_paragraph("• Empieza con 'Quiz del tema' (Título 2)")
    doc.add_paragraph("• Numera las preguntas: 1., 2., 3...")
    doc.add_paragraph("• Las opciones: A., B., C., D.")
    doc.add_paragraph("• Indica la correcta con 'Correcta: X'")
    doc.add_paragraph("• Explicación: opcional pero recomendada")

    doc.add_paragraph().add_run("Asistente IA en la app").bold = True
    doc.add_paragraph("Una vez subido el Word a la app, podrás:")
    doc.add_paragraph("• Generar etiquetas/tags automáticamente")
    doc.add_paragraph("• Crear quizzes mixtos (test + V/F + huecos)")
    doc.add_paragraph("• Sugerir alt-text de imágenes con Claude Vision")
    doc.add_paragraph("• Validar WCAG 2.1 AA con un clic")
    doc.add_paragraph("• Vista previa del SCORM sin descargar")
    doc.add_paragraph("• Detectar problemas de copyright en imágenes")
    doc.add_paragraph("• Reescribir párrafos como callouts automáticamente")

    doc.save(str(output_path))
    return output_path
